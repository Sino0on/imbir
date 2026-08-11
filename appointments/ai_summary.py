"""AI-обработка завершённой консультации: расшифровка записей врача и пациента
по отдельности → резюме (текст + docx) → отправка в чат.

Врач и пациент пишутся отдельными аудиодорожками (см. livekit_integration.services),
поэтому роль говорящего для каждой фразы известна точно — расшифровка каждой
дорожки делается отдельно и объединяется в один хронологический диалог по
таймкодам сегментов, а не угадывается моделью по смыслу.

Вызывается из tasks.generate_consultation_summary после того, как обе дорожки
(насколько они вообще писались) дошли до терминального состояния записи.
"""
import io
import logging
import re

from django.conf import settings
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.utils import timezone
from docx import Document
from openai import OpenAI

from livekit_integration.services import fetch_recording_bytes

logger = logging.getLogger(__name__)

# Лимит Whisper API на размер файла.
MAX_TRANSCRIBE_BYTES = 25 * 1024 * 1024

_ROLES = ('doctor', 'patient')
_ROLE_LABELS = {'doctor': 'Врач', 'patient': 'Пациент'}

SUMMARY_SYSTEM_PROMPT = (
    'Ты медицинский ассистент. Тебе дана расшифровка видео-консультации врача и пациента — '
    'реплики размечены по ролям и времени. Составь краткое структурированное резюме на '
    'русском языке по разделам: "Жалобы пациента", "Что обсудили", "Рекомендации врача", '
    '"Дальнейшие шаги". Опирайся только на разметку ролей из расшифровки, не путай, кто '
    'что сказал. Опирайся только на то, что реально прозвучало в разговоре, ничего не '
    'придумывай. Если содержимого недостаточно для содержательного резюме — так и напиши '
    'одной строкой.'
)


def _transcribe_role(appointment, role: str) -> list[dict]:
    """Транскрибирует запись одной роли с таймкодами сегментов.

    Пустой список, если для роли не было записи (участник не публиковал микрофон
    или запись закончилась ошибкой) — это не сбой, просто той стороне нечего сказать.
    """
    if not getattr(appointment, f'{role}_recording_url'):
        return []

    content, filename = fetch_recording_bytes(appointment, role)
    if len(content) > MAX_TRANSCRIBE_BYTES:
        raise ValueError(
            f'consultation={appointment.id}: запись роли {role} слишком большая для Whisper '
            f'({len(content)} байт, лимит {MAX_TRANSCRIBE_BYTES})'
        )

    file_obj = io.BytesIO(content)
    file_obj.name = filename or f'consultation-{appointment.id}-{role}.ogg'

    client = OpenAI(api_key=settings.OPENAI_API_KEY)
    transcript = client.audio.transcriptions.create(
        model='whisper-1',
        file=file_obj,
        language='ru',
        response_format='verbose_json',
        timestamp_granularities=['segment'],
    )
    return [
        {'start': seg.start, 'text': seg.text.strip()}
        for seg in (transcript.segments or [])
        if seg.text.strip()
    ]


def _format_timestamp(seconds: float) -> str:
    minutes, secs = divmod(int(seconds), 60)
    return f'{minutes:02d}:{secs:02d}'


def build_role_labeled_transcript(appointment) -> str:
    """Расшифровка обеих ролей по отдельности, слитая в один хронологический
    диалог по таймкодам: "[MM:SS] Врач: ...", "[MM:SS] Пациент: ...".
    """
    timeline = []
    for role in _ROLES:
        label = _ROLE_LABELS[role]
        for seg in _transcribe_role(appointment, role):
            timeline.append((seg['start'], f"[{_format_timestamp(seg['start'])}] {label}: {seg['text']}"))

    timeline.sort(key=lambda item: item[0])
    return '\n'.join(line for _, line in timeline)


def summarize_transcript(transcript: str) -> str:
    client = OpenAI(api_key=settings.OPENAI_API_KEY)
    completion = client.chat.completions.create(
        model='gpt-4o-mini',
        messages=[
            {'role': 'system', 'content': SUMMARY_SYSTEM_PROMPT},
            {'role': 'user', 'content': transcript},
        ],
    )
    return completion.choices[0].message.content.strip()


# GPT оформляет разделы как "**Заголовок:** текст" на одной строке —
# в docx превращаем это в заголовок + отдельный абзац.
_SECTION_HEADER_RE = re.compile(r'^\*\*(.+?)\*\*:?\s*(.*)$')


def build_summary_docx(appointment, summary_text: str) -> bytes:
    doc = Document()
    doc.add_heading(f'Резюме консультации №{appointment.id}', level=1)

    doctor_name = appointment.doctor.user.full_name if appointment.doctor else ''
    patient_name = appointment.patient.full_name if appointment.patient else appointment.guest_name
    doc.add_paragraph(f'Дата: {appointment.date} {appointment.time}')
    doc.add_paragraph(f'Врач: {doctor_name}')
    doc.add_paragraph(f'Пациент: {patient_name}')
    doc.add_paragraph('')

    for line in summary_text.splitlines():
        line = line.strip()
        if not line:
            continue
        match = _SECTION_HEADER_RE.match(line)
        if match:
            heading, rest = match.groups()
            doc.add_heading(heading.strip(), level=2)
            if rest.strip():
                doc.add_paragraph(rest.strip())
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def save_summary_docx(appointment, docx_bytes: bytes) -> str:
    """Сохраняет docx через стандартный storage (тот же, что и /api/upload/)
    и возвращает публичную ссылку на скачивание."""
    filename = f'ai-summaries/consultation-{appointment.id}.docx'
    saved_path = default_storage.save(filename, ContentFile(docx_bytes))
    relative_url = settings.MEDIA_URL + saved_path

    if settings.SITE_URL:
        return settings.SITE_URL + relative_url

    logger.warning(
        'AI-резюме consultation=%s: SITE_URL не задан в настройках — ссылка на docx '
        'будет относительной и не откроется напрямую из чата',
        appointment.id,
    )
    return relative_url


def send_summary_to_chat(appointment, summary_text: str, docx_url: str = '') -> None:
    from chat.services import get_or_create_room, send_system_message

    doctor_user = appointment.doctor.user if appointment.doctor else None
    patient_user = appointment.patient
    if not doctor_user or not patient_user:
        logger.warning(
            'AI-резюме consultation=%s: нет врача или пациента, отправлять некуда',
            appointment.id,
        )
        return

    room = get_or_create_room(patient_user, doctor_user)
    content = f'Резюме консультации (сформировано автоматически):\n\n{summary_text}'
    if docx_url:
        content += f'\n\n📄 Скачать в DOCX: {docx_url}'
    send_system_message(room, content)
    logger.info('AI-резюме consultation=%s отправлено в чат room=%s', appointment.id, room.id)


def generate_and_deliver_summary(appointment) -> None:
    logger.info('AI-резюме consultation=%s: начинаем расшифровку записей', appointment.id)
    transcript = build_role_labeled_transcript(appointment)
    if not transcript.strip():
        logger.warning(
            'AI-резюме consultation=%s: расшифровка пустая (обе стороны молчали?), резюме не строим',
            appointment.id,
        )
        return
    logger.info(
        'AI-резюме consultation=%s: расшифровка готова (%d симв.)',
        appointment.id, len(transcript),
    )

    summary = summarize_transcript(transcript)
    logger.info('AI-резюме consultation=%s: резюме сформировано', appointment.id)

    docx_bytes = build_summary_docx(appointment, summary)
    docx_url = save_summary_docx(appointment, docx_bytes)

    appointment.ai_summary = summary
    appointment.ai_summary_generated_at = timezone.now()
    appointment.ai_summary_docx_url = docx_url
    appointment.save(update_fields=[
        'ai_summary', 'ai_summary_generated_at', 'ai_summary_docx_url', 'updated_at',
    ])

    send_summary_to_chat(appointment, summary, docx_url)
